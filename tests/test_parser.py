import os
import sys
import uuid
import pytest
from unittest.mock import MagicMock, patch
from fastapi import HTTPException

# Add the apps/api folder to python path for importing modules
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../apps/api")))

from services.parser import extract_text_from_file, parse_resume_text
from models import User, Profile
from routers.profile import confirm_profile

# 1. DOCX Text Extraction Test
def test_docx_extraction(tmp_path):
    """
    Creates a real docx file using python-docx and verifies that text is extracted.
    """
    import docx
    file_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Experience: 3 years at SaaS Platform.")
    doc.save(str(file_path))
    
    extracted_text = extract_text_from_file(str(file_path), "test.docx")
    assert "Jane Doe" in extracted_text
    assert "Software Engineer" in extracted_text
    assert "SaaS Platform" in extracted_text

# 2. PDF Text Extraction Test
@patch("services.parser.pypdf.PdfReader")
def test_pdf_extraction(mock_pdf_reader, tmp_path):
    """
    Mocks pypdf Reader to verify that text is extracted from all pages of a PDF.
    """
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "Jane Doe\nEmail: developer@example.com\nSkills: TypeScript, Next.js"
    
    mock_instance = mock_pdf_reader.return_value
    mock_instance.pages = [mock_page]
    
    # Create a dummy file representing the PDF binary
    file_path = tmp_path / "test.pdf"
    file_path.write_bytes(b"%PDF-1.4 mock content")
    
    extracted_text = extract_text_from_file(str(file_path), "test.pdf")
    
    assert "Jane Doe" in extracted_text
    assert "developer@example.com" in extracted_text
    assert "Next.js" in extracted_text
    mock_pdf_reader.assert_called_once_with(str(file_path))

# 3. Unsupported File Type Test
def test_unsupported_file_rejection(tmp_path):
    """
    Verifies that calling the extractor with an unsupported extension raises an HTTP 400.
    """
    file_path = tmp_path / "test.png"
    file_path.write_bytes(b"fake png binary data")
    
    with pytest.raises(HTTPException) as exc_info:
        extract_text_from_file(str(file_path), "test.png")
        
    assert exc_info.value.status_code == 400
    assert "Unsupported file type" in exc_info.value.detail

# 4. Profile Confirmation Test
def test_profile_confirmation():
    """
    Verifies that confirm_profile route dependency queries the profile, updates
    confirmed status to True, commits to DB, and returns the updated profile.
    """
    mock_db = MagicMock()
    mock_user_id = uuid.uuid4()
    mock_user = User(id=mock_user_id, email="developer@example.com")
    
    mock_profile = Profile(
        user_id=mock_user_id,
        profile_json={"name": "Jane Doe"},
        confirmed=False
    )
    
    # Setup database query mock sequence
    mock_db.query.return_value.filter.return_value.first.return_value = mock_profile
    
    # Trigger confirm_profile route logic
    result = confirm_profile(db=mock_db, current_user=mock_user)
    
    assert result.confirmed is True
    mock_db.commit.assert_called_once()
